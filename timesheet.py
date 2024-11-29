from pytz import timezone
from datetime import datetime
from docx import Document
import holidays
from calendar import monthrange

# Set Singapore timezone
sg_timezone = timezone("Asia/Singapore")
now = datetime.now(sg_timezone)

# Auto update the month and year
year = now.year
month = now.month

# Create a holiday list for Singapore
sg_holidays = holidays.SG(years=year)

# Get the number of days in the current month
num_days = monthrange(year, month)[1]

# Path to your timesheet document
path = "/Users/zhaoyan.x/Documents/Timesheets/"
doc_filename = "timesheet.docx"  # Adjust if needed
doc = Document(path + doc_filename)


def set_cell_text(cell, text):
    """Helper function to set text in a table cell without extra spacing."""
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = 0


def fill_in_month(table):
    set_cell_text(table.rows[5].cells[6], now.strftime("%B") + " " + str(year))


def is_public_holiday(date):
    return date in sg_holidays


def process_work_hours(row, day_num, start_cell_index, end_cell_index):
    """Process and set the work hours or mark holidays/weekends."""
    if day_num <= num_days:
        # Correctly create a timezone-aware datetime object
        date_obj = datetime(year, month, day_num, tzinfo=sg_timezone)
        weekday = date_obj.weekday()  # 0 = Monday, ..., 6 = Sunday

        if is_public_holiday(date_obj):
            set_cell_text(row.cells[start_cell_index], "PH")
            set_cell_text(row.cells[end_cell_index], "PH")
        else:
            if weekday == 5:  # Saturday
                set_cell_text(row.cells[start_cell_index], "Sat")
                set_cell_text(row.cells[end_cell_index], "Sat")
            elif weekday == 6:  # Sunday
                set_cell_text(row.cells[start_cell_index], "Sun")
                set_cell_text(row.cells[end_cell_index], "Sun")
            else:
                # Weekday work hours
                set_cell_text(row.cells[start_cell_index], "0830")
                set_cell_text(
                    row.cells[end_cell_index], "1800" if weekday < 4 else "1730"
                )


def process_table(table):
    """Process the timesheet table and set the work hours for each row."""
    rows = table.rows

    for i, row in enumerate(rows):
        # Skip rows with no date or invalid date cells
        if not row.cells[0].text.strip().isdigit():
            continue

        day_num_1 = int(
            row.cells[0].text.strip()
        )  # Correctly calculate the day for the row
        day_num_2 = day_num_1 + 15  # Second day for the row (next column block)
        # Process the first set of dates (columns 1-3)
        process_work_hours(row, day_num_1, 1, 3)

        # Process the second set of dates (columns 11-13)
        if day_num_2 <= num_days:
            process_work_hours(row, day_num_2, 11, 13)


# Process the relevant table (assuming the second table is for time logging)
for table_index, table in enumerate(doc.tables):
    if table_index == 0:
        fill_in_month(table)
    if table_index == 1:  # Assuming the 2nd table is for time logging
        process_table(table)

# Save the modified document
doc.save(path + "modified_" + doc_filename)
